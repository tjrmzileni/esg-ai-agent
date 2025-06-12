# Sustainability Reporting AI Agent (MVP)

# --- Setup ---
# Required Libraries
import pandas as pd
import matplotlib.pyplot as plt
import os
from docx import Document
from datetime import datetime
from transformers import pipeline

# --- 1. Load User Data ---
def load_data(filepath):
    if filepath.endswith(".csv"):
        return pd.read_csv(filepath)
    elif filepath.endswith(".xlsx"):
        return pd.read_excel(filepath)
    else:
        raise ValueError("Unsupported file format")

# --- 2. Analyze Data for ESG Metrics ---
def analyze_metrics(df):
    # Placeholder logic – adapt based on ESG columns
    metrics = {}
    metrics['total_emissions'] = df.get('Emissions_tCO2', pd.Series([0])).sum()
    metrics['total_energy'] = df.get('Energy_kWh', pd.Series([0])).sum()
    metrics['total_waste'] = df.get('Waste_kg', pd.Series([0])).sum()
    return metrics

# --- 3. Generate Report ---
def generate_report(metrics, company_name="DemoCo"):
    doc = Document()
    doc.add_heading(f"Sustainability Report – {company_name}", 0)
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")

    doc.add_heading("Key ESG Metrics", level=1)
    for k, v in metrics.items():
        doc.add_paragraph(f"{k.replace('_', ' ').capitalize()}: {v:,}")

    doc.add_heading("Analysis Summary", level=1)
    summarizer = pipeline("summarization")
    summary_input = f"This quarter, {company_name} produced {metrics['total_emissions']} tCO2 emissions, consumed {metrics['total_energy']} kWh energy, and generated {metrics['total_waste']} kg waste."
    summary = summarizer(summary_input, max_length=100, min_length=30, do_sample=False)[0]['summary_text']
    doc.add_paragraph(summary)

    filename = f"sustainability_report_{company_name.lower()}.docx"
    doc.save(filename)
    print(f"Report saved as {filename}")
    return filename

# --- 4. Optional: Visualization ---
def plot_metrics(metrics):
    labels = list(metrics.keys())
    values = list(metrics.values())
    plt.bar(labels, values, color='green')
    plt.title("ESG Metrics Overview")
    plt.ylabel("Amount")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig("esg_metrics_chart.png")
    print("Chart saved as esg_metrics_chart.png")

# --- Example Workflow ---
if __name__ == "__main__":
    df = load_data("ESG.csv")
    metrics = analyze_metrics(df)
    plot_metrics(metrics)
    generate_report(metrics, company_name="GreenFuture Inc")
